import g4f
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from PyQt5.QtWidgets import (QApplication, QWidget, QLabel, QLineEdit, QPushButton,
                             QVBoxLayout, QHBoxLayout, QComboBox, QTextEdit, QFileDialog,
                             QMessageBox, QProgressBar, QDialog)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
import sys


class EssayGeneratorThread(QThread):
    """Поток для генерации реферата."""
    planGenerated = pyqtSignal(str)
    contentGenerated = pyqtSignal(str, str)
    generationFinished = pyqtSignal(bool, str)
    progressUpdated = pyqtSignal(int)

    def __init__(self, topic, provider, model, save_path):
        super().__init__()
        self.topic = topic
        self.provider = provider
        self.model = model
        self.save_path = save_path
        self.essay_generator = EssayGenerator(provider=self.provider, model=self.model)
        self.docx_writer = DocxWriter(self.save_path)
        self.essay = Essay(self.topic, self.essay_generator, self.docx_writer)

    def run(self):
        """Запускает процесс генерации."""
        try:
            success = self.essay.generate()
            if success:
                self.progressUpdated.emit(50)  # 50% после генерации плана и контента
                self.planGenerated.emit('\n'.join(self.essay.plan))

                for plan_item, item_content in zip(self.essay.plan, self.essay.content):
                    self.contentGenerated.emit(plan_item, item_content)

                self.essay.create_docx()
                self.progressUpdated.emit(100) # 100% после сохранения
                self.generationFinished.emit(True, "Реферат успешно сгенерирован и сохранен!")
            else:
                self.generationFinished.emit(False, "Ошибка при генерации реферата.")
        except Exception as e:
            self.generationFinished.emit(False, f"Произошла непредвиденная ошибка: {e}")

class EssayGenerator:
    def __init__(self, provider=g4f.Provider.Copilot, model=g4f.models.gpt_4):
        self.provider = provider
        self.model = model

    def generate_plan(self, topic):
        """Генерирует план реферата с помощью GPT."""
        try:
            response = g4f.ChatCompletion.create(
                model=self.model,
                provider=self.provider,
                messages=[
                    {"role": "system", "content": "Ты должен написать план реферата. И ответ должен только быть в виде нумеровонного списка пунктов план, без подробностей."},
                    {"role": "user", "content": f"Составь краткий план реферата из 8 пунктов на тему: {topic}"},
                ],
            )
            return response
        except Exception as e:
            print(f"Ошибка при генерации плана: {e}")
            return None

    def generate_content(self, plan_item, plan):
        """Генерирует текст для пункта плана, сохраняя контекст."""
        try:
            response = g4f.ChatCompletion.create(
                model=self.model,
                provider=self.provider,
                messages=[
                    {"role": "system", "content": "Ты - опытный автор научных текстов. Напиши подробный текст для данного пункта плана реферата, учитывая предыдущий контекст."},
                    {"role": "user", "content": f"Напиши Пункт плана реферата: {plan_item} из Плана {plan}. Но в ответе не пиши назв пункта, только содержание ничего лишнего"},
                ],
            )
            return response
        except Exception as e:
            print(f"Ошибка при генерации контента: {e}")
            return None

class DocxWriter:
    def __init__(self, filename):
        self.document = Document()
        self.filename = filename
        self.setup_styles()

    def setup_styles(self):
        """Настраивает стили для документа."""
        styles = self.document.styles

        # Стиль заголовка 1
        heading1_style = styles.add_style('Heading1Custom', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.base_style = styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Times New Roman'
        heading1_font.size = Pt(16)
        heading1_font.bold = True

        # Стиль обычного текста
        normal_style = styles.add_style('NormalCustom', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.base_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(14)

    def add_heading(self, text, level=1):
        """Добавляет заголовок в документ."""
        self.document.add_heading(text, level=level)

    def add_paragraph(self, text, style='NormalCustom'):
        """Добавляет параграф в документ."""
        self.document.add_paragraph(text, style=style)

    def add_page_break(self):
        """Добавляет разрыв страницы."""
        self.document.add_page_break()

    def save(self):
        """Сохраняет документ."""
        self.document.save(self.filename)

class Essay:
    def __init__(self, topic, generator, writer):
        self.topic = topic
        self.generator = generator
        self.writer = writer
        self.plan = None
        self.content = []

    def generate(self):
        """Генерирует план и контент реферата."""
        self.plan = self.generator.generate_plan(self.topic)
        if self.plan:
            print("Сгенерированный план:")
            print(self.plan)
            self.plan = [item.strip() for item in self.plan.split('\n') if item.strip()]
            for plan_item in self.plan:
                item_content = self.generator.generate_content(plan_item, self.plan)
                if item_content:
                    self.content.append(item_content)
                    print(f'добавлен {plan_item}')
                else:
                    print("Ошибка при генерации контента для пункта плана.")
                    return False
            return True
        else:
            print("Не удалось сгенерировать план.")
            return False

    def create_docx(self):
        """Создает и сохраняет docx файл с рефератом."""
        self.writer.add_heading(f"Реферат на тему: {self.topic}", level=1)
        self.writer.add_page_break()
        self.writer.add_paragraph("План", style='Heading1Custom')

        for plan_item in self.plan:
            self.writer.add_paragraph(plan_item, style='Heading1Custom')
        self.writer.add_page_break()

        for plan_item, item_content in zip(self.plan, self.content):
            self.writer.add_paragraph(plan_item, style='Heading1Custom')
            self.writer.add_paragraph(item_content, style='NormalCustom')

        self.writer.save()
        print(f"Реферат сохранен в файл {self.writer.filename}")


class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Генератор рефератов")
        self.setGeometry(300, 300, 600, 400)

        self.initUI()

    def initUI(self):
        layout = QVBoxLayout()

        # Тема реферата
        topic_layout = QHBoxLayout()
        self.topic_label = QLabel("Тема реферата:")
        self.topic_input = QLineEdit()
        topic_layout.addWidget(self.topic_label)
        topic_layout.addWidget(self.topic_input)
        layout.addLayout(topic_layout)

        # Провайдер
        provider_layout = QHBoxLayout()
        self.provider_label = QLabel("Провайдер:")
        self.provider_combo = QComboBox()
        self.provider_combo.addItems([provider.__name__ for provider in g4f.Provider.__providers__ if provider.working])
        provider_layout.addWidget(self.provider_label)
        provider_layout.addWidget(self.provider_combo)
        layout.addLayout(provider_layout)

        # Модель
        model_layout = QHBoxLayout()
        self.model_label = QLabel("Модель:")
        self.model_combo = QComboBox()
        self.model_combo.addItem("gpt-4")
        self.model_combo.addItem("gpt_35_long")
        self.model_combo.addItem("gpt-3.5-turbo")
        model_layout.addWidget(self.model_label)
        model_layout.addWidget(self.model_combo)
        layout.addLayout(model_layout)

        # Кнопка "Сгенерировать"
        self.generate_button = QPushButton("Сгенерировать")
        self.generate_button.clicked.connect(self.generate_essay)
        layout.addWidget(self.generate_button)

        # Прогресс бар
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)

        # План
        self.plan_label = QLabel("План:")
        self.plan_text = QTextEdit()
        self.plan_text.setReadOnly(True)
        layout.addWidget(self.plan_label)
        layout.addWidget(self.plan_text)

        # Содержание
        self.content_label = QLabel("Содержание:")
        self.content_text = QTextEdit()
        self.content_text.setReadOnly(True)
        layout.addWidget(self.content_label)
        layout.addWidget(self.content_text)

        self.setLayout(layout)

    def generate_essay(self):
        """Запускает генерацию реферата в отдельном потоке."""
        topic = self.topic_input.text().strip()
        if not topic:
            QMessageBox.warning(self, "Ошибка", "Введите тему реферата!")
            return

        provider_name = self.provider_combo.currentText()
        provider = getattr(g4f.Provider, provider_name)

        model_name = self.model_combo.currentText()
        if model_name == 'gpt-4':
            model = g4f.models.gpt_4
        elif model_name == 'gpt_35_long':
            model = g4f.models.gpt_35_long
        else:
            model = g4f.models.default

        file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить реферат", "", "Word Documents (*.docx)")
        if not file_path:
            return
        if not file_path.endswith(".docx"):
            file_path += ".docx"

        self.progress_bar.setValue(0)
        self.plan_text.clear()
        self.content_text.clear()
        self.generate_button.setEnabled(False)

        self.essay_thread = EssayGeneratorThread(topic, provider, model, file_path)
        self.essay_thread.planGenerated.connect(self.update_plan)
        self.essay_thread.contentGenerated.connect(self.update_content)
        self.essay_thread.generationFinished.connect(self.on_generation_finished)
        self.essay_thread.progressUpdated.connect(self.update_progress)
        self.essay_thread.start()

    def update_plan(self, plan):
        self.plan_text.setText(plan)

    def update_content(self, plan_item, content):
        self.content_text.append(f"{plan_item}\n{content}\n")

    def on_generation_finished(self, success, message):
        self.generate_button.setEnabled(True)
        if success:
            QMessageBox.information(self, "Готово", message)
        else:
            QMessageBox.warning(self, "Ошибка", message)

    def update_progress(self, value):
        self.progress_bar.setValue(value)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())